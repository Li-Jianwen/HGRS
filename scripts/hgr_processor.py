#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HGR Processor
下载PDF/DOCX → 提取表格 → 分类整理
"""

import os
import sys
import io
import logging
import re
import requests
import pdfplumber
from typing import List, Dict, Optional, Tuple

sys.path.insert(0, os.path.dirname(__file__))
from excel_writer import classify_by_approval_no

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

class HGRProcessor:
    def __init__(self, config: Dict, logger=None):
        self.config = config
        self.logger = logger or logging.getLogger(__name__)
        self.retry_count = int(config.get('download_retry', 3))
        self.timeout = int(config.get('download_timeout', 30))

    def download_file(self, url: str, cookies: Dict, save_path: str) -> bool:
        """
        下载文件到本地（支持PDF和DOCX）
        """
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Referer': 'https://apply.hgrg.net/',
        }
        
        for attempt in range(self.retry_count):
            try:
                r = requests.get(url, cookies=cookies, headers=headers, timeout=self.timeout)
                if r.status_code == 200:
                    with open(save_path, 'wb') as f:
                        f.write(r.content)
                    size_kb = len(r.content) // 1024
                    content_type = r.headers.get('Content-Type', '未知')
                    self.logger.info(f"✅ 下载成功 ({content_type}, {size_kb}KB): {url}")
                    return True
                else:
                    self.logger.warning(f"⚠️ 下载失败 ({attempt+1}/{self.retry_count}): HTTP {r.status_code}")
            except Exception as e:
                self.logger.warning(f"⚠️ 下载异常 ({attempt+1}/{self.retry_count}): {str(e)}")
        
        self.logger.error(f"❌ 下载失败，已重试{self.retry_count}次: {url}")
        return False

    @staticmethod
    def clean_cell(cell: Optional[str]) -> str:
        """
        清洗单元格内容：替换换行、去除多余空格
        """
        if not cell:
            return ''
        # PDFplumber的单元格可能有换行，需要替换成空格
        cleaned = str(cell).replace('\n', ' ').replace('\r', ' ').strip()
        # 合并连续空格
        cleaned = re.sub(r'\s+', ' ', cleaned)
        return cleaned

    @staticmethod
    def classify_by_approval_no(approval_no: str) -> str:
        return classify_by_approval_no(approval_no)

    def extract_tables_from_pdf(self, pdf_path: str) -> Dict[str, List[Dict]]:
        """
        从PDF提取所有表格数据，按审批类型分类
        返回: { 类型: [ { 列名: 值 }, ... ] }
        """
        result = {
            '采集审批': [],
            '保藏审批': [],
            '国际科学研究合作审批': [],
            '材料出境证明': [],
        }

        # 表头列定义
        col_defs = {
            '采集审批': ['序号', '审批号', '项目名称', '申请单位', '批准时间'],
            '保藏审批': ['序号', '审批号', '项目名称', '申请单位', '批准时间'],
            '国际科学研究合作审批': ['序号', '审批号', '项目名称', '医疗机构(组长单位)', '申办方', '合同研究组织', '检测/数据单位', '批准时间'],
            '材料出境证明': ['序号', '审批号', '项目名称', '申请单位', '批准时间'],
        }

        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_idx, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    if not tables:
                        continue
                    
                    for table in tables:
                        # 跳过只有表头的空表格
                        has_data = False
                        for row in table[1:]:
                            if any(self.clean_cell(cell) for cell in row):
                                has_data = True
                                break
                        if not has_data:
                            continue

                        # 处理每一行数据
                        for row_idx, row in enumerate(table):
                            # Clean all cells
                            cleaned_row = [self.clean_cell(cell) for cell in row]
                            
                            # Skip header rows or empty rows
                            if not cleaned_row[0]:
                                continue
                            if '序号' in cleaned_row[0] and '审批号' in (cleaned_row[1] if len(cleaned_row) > 1 else ''):
                                continue  # 表头跳过
                            if '出口出境证明' in cleaned_row[0]:
                                continue  # 小节标题跳过

                            # 获取审批号
                            if len(cleaned_row) < 2:
                                self.logger.warning(f"⚠️ 跳过短行 (只有 {len(cleaned_row)} 列): {cleaned_row}")
                                continue
                            
                            approval_no = cleaned_row[1]
                            category = self.classify_by_approval_no(approval_no)
                            expected_cols = col_defs[category]
                            
                            # 填充到dict
                            row_dict = {'批次': None}  # 预留批次字段
                            for col_idx, col_name in enumerate(expected_cols):
                                if col_idx < len(cleaned_row):
                                    row_dict[col_name] = cleaned_row[col_idx]
                                else:
                                    row_dict[col_name] = ''
                            
                            result[category].append(row_dict)
                            self.logger.debug(f"  ✓ [{category}] 序号={row_dict.get('序号', '')} 审批号={approval_no}")

                # end for tables
            # end with pdf
        except Exception as e:
            self.logger.error(f"❌ PDF提取失败 {pdf_path}: {str(e)}", exc_info=True)
            return None

        # 统计
        total = sum(len(rows) for rows in result.values())
        self.logger.info(f"✅ PDF提取完成: 总计 {total} 条记录")
        for cat, rows in result.items():
            if rows:
                self.logger.info(f"   {cat}: {len(rows)} 条")
        
        return result

    def extract_tables_from_docx(self, docx_path: str) -> Dict[str, List[Dict]]:
        """
        从DOCX提取所有表格数据，按审批类型分类
        返回: { 类型: [ { 列名: 值 }, ... ] }
        """
        result = {
            '采集审批': [],
            '保藏审批': [],
            '国际科学研究合作审批': [],
            '材料出境证明': [],
        }

        # 表头列定义（与PDF版本一致）
        col_defs = {
            '采集审批': ['序号', '审批号', '项目名称', '申请单位', '批准时间'],
            '保藏审批': ['序号', '审批号', '项目名称', '申请单位', '批准时间'],
            '国际科学研究合作审批': ['序号', '审批号', '项目名称', '医疗机构(组长单位)', '申办方', '合同研究组织', '检测/数据单位', '批准时间'],
            '材料出境证明': ['序号', '审批号', '项目名称', '申请单位', '批准时间'],
        }

        try:
            from docx import Document
            doc = Document(docx_path)
            
            if not doc.tables:
                self.logger.warning(f"⚠️ DOCX中未找到表格: {docx_path}")
                return result

            for table in doc.tables:
                rows = table.rows
                if len(rows) < 2:
                    continue  # 至少需表头+一行数据

                # 逐行处理（跳过表头）
                data_start = 0
                for row_idx, row in enumerate(rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    # 跳过空行
                    if not any(cells):
                        continue
                    
                    # 检测表头行
                    first_cell = cells[0] if cells else ''
                    if ('序号' in first_cell or '序\n号' in first_cell or 
                        (row_idx == 0 and not first_cell.isdigit())):
                        data_start = row_idx + 1
                        continue
                    if '出口出境证明' in first_cell:
                        continue

                    # 处理数据行
                    if row_idx < data_start:
                        continue

                    # 清洗单元格
                    cleaned = [self.clean_cell(c) for c in cells]
                    
                    if len(cleaned) < 2 or not cleaned[0]:
                        continue

                    approval_no = cleaned[1]
                    category = self.classify_by_approval_no(approval_no)
                    expected_cols = col_defs[category]
                    
                    row_dict = {'批次': None}
                    for col_idx, col_name in enumerate(expected_cols):
                        if col_idx < len(cleaned):
                            row_dict[col_name] = cleaned[col_idx]
                        else:
                            row_dict[col_name] = ''
                    
                    result[category].append(row_dict)
                    self.logger.debug(f"  ✓ [{category}] 序号={row_dict.get('序号', '')} 审批号={approval_no}")

        except ImportError:
            self.logger.error("❌ 需要安装 python-docx: pip install python-docx")
            return None
        except Exception as e:
            self.logger.error(f"❌ DOCX提取失败 {docx_path}: {str(e)}", exc_info=True)
            return None

        # 统计
        total = sum(len(rows) for rows in result.values())
        self.logger.info(f"✅ DOCX提取完成: 总计 {total} 条记录")
        for cat, rows in result.items():
            if rows:
                self.logger.info(f"   {cat}: {len(rows)} 条")
        
        return result

    def process_pdf(self, pdf_url: str, cookies: Dict, save_dir: str, 
                   year: int, batch: int, title: str) -> Optional[Dict]:
        """
        完整处理一个批次：下载文件 → 提取表格 → 返回结构化数据
        支持PDF和DOCX两种格式
        """
        # 根据URL扩展名确定文件类型
        ext = os.path.splitext(pdf_url)[1].lower()
        if ext not in ('.pdf', '.docx'):
            ext = '.pdf'  # 默认PDF
        
        file_name = f"{year}_{batch}{ext}"
        file_path = os.path.join(save_dir, file_name)
        os.makedirs(save_dir, exist_ok=True)

        # 下载
        if not self.download_file(pdf_url, cookies, file_path):
            return None

        # 根据格式提取
        if ext == '.docx':
            data = self.extract_tables_from_docx(file_path)
        else:
            data = self.extract_tables_from_pdf(file_path)
        
        if data is None:
            return None

        # 添加批次信息到每条记录
        for category in data:
            for row in data[category]:
                row['批次'] = f"{year}年第{batch}批"
                row['批次_year'] = year
                row['批次_batch'] = batch

        return {
            'year': year,
            'batch': batch,
            'title': title,
            'pdf_path': file_path,
            'data': data,
            'total_count': sum(len(rows) for rows in data.values()),
        }


if __name__ == '__main__':
    # 测试入口
    import argparse
    import json
    
    logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
    
    parser = argparse.ArgumentParser()
    parser.add_argument('--pdf-url', required=True)
    parser.add_argument('--output-dir', required=True)
    parser.add_argument('--year', type=int, required=True)
    parser.add_argument('--batch', type=int, required=True)
    parser.add_argument('--title', default='')
    args = parser.parse_args()
    
    # Dummy config
    processor = HGRProcessor({'download_retry': 3, 'download_timeout': 30})
    
    # Cookies从文件读取
    if os.path.exists('cookies.json'):
        with open('cookies.json', 'r', encoding='utf-8') as f:
            cookies = json.load(f)
    else:
        cookies = {}
    
    result = processor.process_pdf(args.pdf_url, cookies, args.output_dir, args.year, args.batch, args.title)
    if result:
        print(f"\n=== 处理完成 ===")
        print(f"批次: {args.year}年第{args.batch}批")
        print(f"总记录数: {result['total_count']}")
        for cat, rows in result['data'].items():
            if rows:
                print(f"  {cat}: {len(rows)}")
